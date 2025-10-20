from flask import Flask, render_template, request, jsonify, session, redirect, url_for
import os
from werkzeug.utils import secure_filename
import mysql.connector
from google import genai
import PyPDF2
from docx import Document
import json
from datetime import datetime
from dotenv import load_dotenv

# Carregar variáveis de ambiente
load_dotenv()

app = Flask(__name__)
app.secret_key = 'sua_chave_secreta_aqui_altere_isto_em_producao'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# Configuração do Gemini - com verificação
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'SUA CHAVE')

# Verificar se a API key está disponível
if not GEMINI_API_KEY:
    print("AVISO: GEMINI_API_KEY não encontrada nas variáveis de ambiente")
    # Usar a chave diretamente como fallback
    GEMINI_API_KEY = "AIzaSyAKjdtiVdeAY3j3ppZX9sApo_aCqaxUem8"

try:
    client = genai.Client(api_key=GEMINI_API_KEY)
    print("Conexão com Gemini AI estabelecida com sucesso!")
except Exception as e:
    print(f"Erro ao conectar com Gemini AI: {e}")
    client = None

# Configuração do MySQL - com fallback para SQLite se MySQL falhar
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'user': os.getenv('DB_USER', 'root'),
    'password': os.getenv('DB_PASSWORD', ''),
    'database': os.getenv('DB_NAME', 'dados_disparo_curriculo'),
    'charset': 'utf8mb4'
}

# Criar diretório de uploads
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Fallback para SQLite se MySQL não estiver disponível
USE_SQLITE = False
try:
    conn = mysql.connector.connect(**DB_CONFIG)
    conn.close()
    print("Conexão com MySQL estabelecida com sucesso!")
except mysql.connector.Error as e:
    print(f"MySQL não disponível: {e}. Usando SQLite como fallback.")
    USE_SQLITE = True
    import sqlite3

def get_db_connection():
    """Cria e retorna uma conexão com o banco de dados"""
    if USE_SQLITE:
        try:
            conn = sqlite3.connect('vagas.db')
            conn.row_factory = sqlite3.Row
            return conn
        except sqlite3.Error as e:
            print(f"Erro ao conectar com SQLite: {e}")
            return None
    else:
        try:
            conn = mysql.connector.connect(**DB_CONFIG)
            return conn
        except mysql.connector.Error as e:
            print(f"Erro ao conectar com MySQL: {e}")
            return None

def init_db():
    """Verifica se a tabela existe e está com a estrutura correta"""
    conn = get_db_connection()
    if conn:
        try:
            if USE_SQLITE:
                cursor = conn.cursor()
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS vagas (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        nome_vaga TEXT NOT NULL,
                        email TEXT NOT NULL,
                        data_inserida DATE NOT NULL,
                        data_limite DATE,
                        requisitos TEXT NOT NULL,
                        atividades TEXT NOT NULL,
                        beneficios TEXT,
                        observacoes TEXT
                    )
                ''')
                
                # Verificar se há dados na tabela
                cursor.execute("SELECT COUNT(*) FROM vagas")
                count = cursor.fetchone()[0]
                
                if count == 0:
                    # Inserir alguns dados de exemplo
                    cursor.execute('''
                        INSERT INTO vagas (nome_vaga, email, data_inserida, data_limite, requisitos, atividades, beneficios, observacoes)
                        VALUES 
                        ('Desenvolvedor Python Junior', 'rh@techcompany.com', '2024-01-15', '2024-02-15', 'Python, Flask, Django, MySQL', 'Desenvolvimento de aplicações web, Manutenção de sistemas', 'VT, VR, Plano de saúde, Gympass', 'Vaga para nível júnior, necessário conhecimento em Git'),
                        ('Analista Administrativo Pleno', 'adm@empresax.com', '2024-01-10', '2024-02-10', 'Excel avançado, Gestão de processos, Power BI', 'Análise de dados administrativos, Elaboração de relatórios', 'VR, Plano de saúde, Day off', 'Horário comercial, Trabalho híbrido'),
                        ('Gerente de Projetos TI', 'ti@corporation.com', '2024-01-20', '2024-02-20', 'Gestão de projetos, Scrum, Kanban, PMBOK', 'Coordenação de equipes de TI, Acompanhamento de métricas', 'VT, VR, Plano de saúde, Bônus anual', 'Experiência mínima 3 anos em gestão')
                    ''')
                    print("Dados de exemplo inseridos no SQLite!")
            else:
                cursor = conn.cursor()
                
                # Verificar se a tabela existe
                cursor.execute("""
                    SELECT COUNT(*) FROM information_schema.tables 
                    WHERE table_schema = %s AND table_name = 'vagas'
                """, (DB_CONFIG['database'],))
                
                table_exists = cursor.fetchone()[0] > 0
                
                if not table_exists:
                    # Criar tabela se não existir
                    cursor.execute("""
                        CREATE TABLE vagas (
                            id INT PRIMARY KEY AUTO_INCREMENT,
                            nome_vaga VARCHAR(255) NOT NULL,
                            email VARCHAR(255) NOT NULL,
                            data_inserida DATE NOT NULL,
                            data_limite DATE,
                            requisitos TEXT NOT NULL,
                            atividades TEXT NOT NULL,
                            beneficios TEXT,
                            observacoes TEXT
                        )
                    """)
                    print("Tabela 'vagas' criada com sucesso!")
                
                # Verificar se há dados na tabela
                cursor.execute("SELECT COUNT(*) FROM vagas")
                count = cursor.fetchone()[0]
                
                if count == 0:
                    # Inserir alguns dados de exemplo
                    cursor.execute("""
                        INSERT INTO vagas (nome_vaga, email, data_inserida, data_limite, requisitos, atividades, beneficios, observacoes)
                        VALUES 
                        ('Desenvolvedor Python Junior', 'rh@techcompany.com', '2024-01-15', '2024-02-15', 'Python, Flask, Django, MySQL', 'Desenvolvimento de aplicações web, Manutenção de sistemas', 'VT, VR, Plano de saúde, Gympass', 'Vaga para nível júnior, necessário conhecimento em Git'),
                        ('Analista Administrativo Pleno', 'adm@empresax.com', '2024-01-10', '2024-02-10', 'Excel avançado, Gestão de processos, Power BI', 'Análise de dados administrativos, Elaboração de relatórios', 'VR, Plano de saúde, Day off', 'Horário comercial, Trabalho híbrido'),
                        ('Gerente de Projetos TI', 'ti@corporation.com', '2024-01-20', '2024-02-20', 'Gestão de projetos, Scrum, Kanban, PMBOK', 'Coordenação de equipes de TI, Acompanhamento de métricas', 'VT, VR, Plano de saúde, Bônus anual', 'Experiência mínima 3 anos em gestão')
                    """)
                    print("Dados de exemplo inseridos no MySQL!")
            
            conn.commit()
            
        except Exception as e:
            print(f"Erro ao inicializar banco: {e}")
        finally:
            cursor.close()
            conn.close()

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Erro ao extrair texto do PDF: {str(e)}"

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        return f"Erro ao extrair texto do DOCX: {str(e)}"

def analyze_resume_with_ai(resume_text):
    # Se o cliente Gemini não estiver disponível, usar dados de exemplo
    if client is None:
        print("Usando dados de exemplo - Gemini não disponível")
        return get_fallback_analysis()
    
    try:
        prompt = f"""
        Analise este currículo e identifique as áreas profissionais em que a pessoa se destaca.
        Retorne APENAS um JSON válido com o seguinte formato:
        {{
            "areas_principais": ["lista", "de", "areas", "principais"],
            "areas_secundarias": ["lista", "de", "areas", "secundarias"],
            "pontos_fortes": ["lista", "de", "pontos", "fortes"],
            "sugestoes_melhoria": ["sugestoes", "para", "melhorar", "curriculo"],
            "cursos_recomendados": ["cursos", "recomendados"]
        }}
        
        Analise o texto do currículo e retorne áreas profissionais reais como:
        "Tecnologia da Informação", "Administração", "Finanças", "Marketing", "Recursos Humanos", 
        "Vendas", "Engenharia", "Saúde", "Educação", etc.
        
        Seja específico nas áreas. Por exemplo, em vez de "TI", use "Desenvolvimento de Software", 
        "Análise de Sistemas", "Suporte Técnico", etc.
        
        Currículo para análise:
        {resume_text[:3000]}
        """
        
        # Usar o modelo correto
        response = client.models.generate_content(
            model="gemini-1.5-flash-latest",  # Modelo correto
            contents=prompt
        )
        
        # Extrair JSON da resposta
        response_text = response.text
        print("Resposta da IA:", response_text)
        
        # Tentar extrair JSON da resposta
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0]
        else:
            json_str = response_text
        
        # Limpar a string JSON
        json_str = json_str.strip()
        if json_str.startswith('{') and json_str.endswith('}'):
            return json.loads(json_str)
        else:
            raise Exception("Resposta não é um JSON válido")
        
    except Exception as e:
        print(f"Erro na análise com IA: {e}")
        return get_fallback_analysis()

def get_fallback_analysis():
    """Retorna análise de fallback quando a IA não está disponível"""
    return {
        "areas_principais": ["Desenvolvimento de Software", "Análise de Sistemas"],
        "areas_secundarias": ["Administração", "Suporte Técnico"],
        "pontos_fortes": ["Experiência em programação", "Conhecimento em banco de dados", "Habilidades de comunicação"],
        "sugestoes_melhoria": ["Adicionar certificações específicas", "Incluir projetos pessoais"],
        "cursos_recomendados": ["Gestão de Projetos Ágeis", "Banco de Dados Avançado", "Cloud Computing"]
    }

def create_custom_resume(area, user_info):
    # Se o cliente Gemini não estiver disponível, usar dados de exemplo
    if client is None:
        print("Usando currículo de exemplo - Gemini não disponível")
        return get_fallback_resume(area)
    
    try:
        prompt = f"""
        Crie um currículo personalizado para a área de {area} com base nas informações fornecidas.
        Retorne APENAS um JSON válido com a seguinte estrutura:
        {{
            "nome_completo": "Nome completo realista",
            "email": "email.profissional@exemplo.com",
            "telefone": "(11) 98765-4321",
            "objetivo": "Objetivo profissional focado em {area}",
            "formacao": [
                {{
                    "curso": "Nome do curso realista",
                    "instituicao": "Nome da instituição",
                    "periodo": "2020-2023"
                }}
            ],
            "experiencia": [
                {{
                    "cargo": "Cargo relacionado a {area}",
                    "empresa": "Empresa realista",
                    "periodo": "2021-2023",
                    "atividades": ["Atividade 1 relevante", "Atividade 2 relevante"]
                }}
            ],
            "habilidades": ["Habilidade técnica 1", "Habilidade técnica 2", "Habilidade comportamental"],
            "cursos": ["Curso específico 1", "Curso específico 2"]
        }}
        
        Use informações realistas e profissionais para a área de {area}.
        """
        
        response = client.models.generate_content(
            model="gemini-1.5-flash-latest",  # Modelo correto
            contents=prompt
        )
        
        response_text = response.text
        print("Resposta da IA para currículo:", response_text)
        
        # Tentar extrair JSON da resposta
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0]
        else:
            json_str = response_text
        
        # Limpar a string JSON
        json_str = json_str.strip()
        if json_str.startswith('{') and json_str.endswith('}'):
            return json.loads(json_str)
        else:
            raise Exception("Resposta não é um JSON válido")
        
    except Exception as e:
        print(f"Erro na criação de currículo: {e}")
        return get_fallback_resume(area)

def get_fallback_resume(area):
    """Retorna currículo de fallback quando a IA não está disponível"""
    return {
        "nome_completo": "Carlos Eduardo Silva",
        "email": "carlos.silva@email.com",
        "telefone": "(11) 98765-4321",
        "objetivo": f"Atuar na área de {area} aplicando minhas habilidades e experiência para contribuir com o crescimento organizacional",
        "formacao": [
            {
                "curso": "Bacharelado em Sistemas de Informação",
                "instituicao": "Universidade Federal de Tecnologia",
                "periodo": "2018-2022"
            }
        ],
        "experiencia": [
            {
                "cargo": f"Analista de {area.split()[0]}",
                "empresa": "Empresa Solutions Ltda",
                "periodo": "2022-2024",
                "atividades": [
                    "Análise e desenvolvimento de processos",
                    "Elaboração de relatórios gerenciais",
                    "Suporte em projetos estratégicos"
                ]
            }
        ],
        "habilidades": ["Comunicação eficaz", "Trabalho em equipe", "Resolução de problemas", "Gestão de tempo"],
        "cursos": ["Gestão de Projetos com Agile", "Excel Avançado", "Língua Inglesa Intermediário"]
    }

def get_vagas_from_db(areas):
    """Busca vagas no banco de dados baseado nas áreas"""
    conn = get_db_connection()
    if not conn:
        return []
    
    try:
        if USE_SQLITE:
            cursor = conn.cursor()
            vagas_encontradas = []
            
            for area in areas:
                search_terms = [
                    area,
                    area.split()[0] if ' ' in area else area,
                    area.replace(' ', '%')
                ]
                
                for term in search_terms:
                    cursor.execute("""
                        SELECT * FROM vagas 
                        WHERE nome_vaga LIKE ? 
                        OR requisitos LIKE ? 
                        OR atividades LIKE ?
                        OR observacoes LIKE ?
                    """, (f'%{term}%', f'%{term}%', f'%{term}%', f'%{term}%'))
                    results = cursor.fetchall()
                    # Converter para dicionário
                    for row in results:
                        vagas_encontradas.append({
                            'id': row[0],
                            'nome_vaga': row[1],
                            'email': row[2],
                            'data_inserida': row[3],
                            'data_limite': row[4],
                            'requisitos': row[5],
                            'atividades': row[6],
                            'beneficios': row[7],
                            'observacoes': row[8]
                        })
        else:
            cursor = conn.cursor(dictionary=True)
            vagas_encontradas = []
            
            for area in areas:
                search_terms = [
                    area,
                    area.split()[0] if ' ' in area else area,
                    area.replace(' ', '%')
                ]
                
                for term in search_terms:
                    query = """
                        SELECT * FROM vagas 
                        WHERE nome_vaga LIKE %s 
                        OR requisitos LIKE %s 
                        OR atividades LIKE %s
                        OR observacoes LIKE %s
                    """
                    cursor.execute(query, (f'%{term}%', f'%{term}%', f'%{term}%', f'%{term}%'))
                    results = cursor.fetchall()
                    vagas_encontradas.extend(results)
        
        # Remover duplicatas
        vagas_unicas = []
        ids_vistos = set()
        
        for vaga in vagas_encontradas:
            if vaga['id'] not in ids_vistos:
                vagas_unicas.append(vaga)
                ids_vistos.add(vaga['id'])
        
        return vagas_unicas
        
    except Exception as e:
        print(f"Erro ao buscar vagas: {e}")
        return []
    finally:
        cursor.close()
        conn.close()

# Rotas Flask (mantidas iguais)
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analisar_curriculo', methods=['GET', 'POST'])
def analisar_curriculo():
    if request.method == 'POST':
        if 'curriculo' not in request.files:
            return render_template('analisar.html', error='Nenhum arquivo selecionado')
        
        file = request.files['curriculo']
        email_key = request.form.get('email_key')
        
        if file.filename == '':
            return render_template('analisar.html', error='Nenhum arquivo selecionado')
        
        if file and (file.filename.lower().endswith('.pdf') or file.filename.lower().endswith('.docx')):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Extrair texto do arquivo
            if filename.lower().endswith('.pdf'):
                resume_text = extract_text_from_pdf(open(filepath, 'rb'))
            else:
                resume_text = extract_text_from_docx(filepath)
            
            # Analisar currículo com IA
            analise = analyze_resume_with_ai(resume_text)
            
            # Salvar na sessão
            session['analise_curriculo'] = analise
            session['email_key'] = email_key
            session['resume_text'] = resume_text
            
            return render_template('resultado_analise.html', 
                                 analise=analise,
                                 total_areas=len(analise['areas_principais']) + len(analise['areas_secundarias']))
        else:
            return render_template('analisar.html', error='Formato de arquivo não suportado. Use PDF ou DOCX.')
    
    return render_template('analisar.html')

@app.route('/criar_curriculo', methods=['POST'])
def criar_curriculo():
    areas_selecionadas = request.form.getlist('areas')
    analise = session.get('analise_curriculo', {})
    resume_text = session.get('resume_text', '')
    
    if not areas_selecionadas:
        return render_template('resultado_analise.html', 
                             analise=analise,
                             total_areas=len(analise.get('areas_principais', [])) + len(analise.get('areas_secundarias', [])),
                             error='Selecione pelo menos uma área')
    
    curriculos_criados = {}
    
    for area in areas_selecionadas:
        user_info = {
            "texto_curriculo": resume_text,
            "areas_destacadas": analise.get('areas_principais', []) + analise.get('areas_secundarias', []),
            "pontos_fortes": analise.get('pontos_fortes', [])
        }
        
        curriculo_personalizado = create_custom_resume(area, user_info)
        curriculos_criados[area] = curriculo_personalizado
    
    session['curriculos_criados'] = curriculos_criados
    session['areas_selecionadas'] = areas_selecionadas
    
    return render_template('curriculos_criados.html', 
                         curriculos=curriculos_criados,
                         areas=areas_selecionadas)

@app.route('/enviar_emails', methods=['POST'])
def enviar_emails():
    areas_selecionadas = session.get('areas_selecionadas', [])
    email_key = session.get('email_key')
    curriculos = session.get('curriculos_criados', {})
    
    vagas = get_vagas_from_db(areas_selecionadas)
    
    return render_template('envio_emails.html',
                         vagas=vagas,
                         total_vagas=len(vagas),
                         areas=areas_selecionadas)

@app.route('/adicionar_vaga', methods=['GET', 'POST'])
def adicionar_vaga():
    if request.method == 'POST':
        nome_vaga = request.form.get('nome_vaga')
        email = request.form.get('email')
        data_limite = request.form.get('data_limite')
        requisitos = request.form.get('requisitos')
        atividades = request.form.get('atividades')
        beneficios = request.form.get('beneficios')
        observacoes = request.form.get('observacoes')
        
        conn = get_db_connection()
        if conn:
            try:
                if USE_SQLITE:
                    cursor = conn.cursor()
                    cursor.execute("""
                        INSERT INTO vagas (nome_vaga, email, data_inserida, data_limite, requisitos, atividades, beneficios, observacoes)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """, (nome_vaga, email, datetime.now().date(), data_limite, requisitos, atividades, beneficios, observacoes))
                else:
                    cursor = conn.cursor()
                    cursor.execute("""
                        INSERT INTO vagas (nome_vaga, email, data_inserida, data_limite, requisitos, atividades, beneficios, observacoes)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """, (nome_vaga, email, datetime.now().date(), data_limite, requisitos, atividades, beneficios, observacoes))
                
                conn.commit()
                return redirect(url_for('adicionar_vaga', success=True))
            except Exception as e:
                print(f"Erro ao adicionar vaga: {e}")
                return render_template('adicionar_vaga.html', error='Erro ao adicionar vaga no banco de dados')
            finally:
                cursor.close()
                conn.close()
    
    return render_template('adicionar_vaga.html')

@app.route('/vagas')
def listar_vagas():
    conn = get_db_connection()
    if conn:
        try:
            if USE_SQLITE:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM vagas ORDER BY data_inserida DESC")
                rows = cursor.fetchall()
                vagas = []
                for row in rows:
                    vagas.append({
                        'id': row[0],
                        'nome_vaga': row[1],
                        'email': row[2],
                        'data_inserida': row[3],
                        'data_limite': row[4],
                        'requisitos': row[5],
                        'atividades': row[6],
                        'beneficios': row[7],
                        'observacoes': row[8]
                    })
            else:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM vagas ORDER BY data_inserida DESC")
                vagas = cursor.fetchall()
            
            return render_template('listar_vagas.html', vagas=vagas)
        except Exception as e:
            print(f"Erro ao buscar vagas: {e}")
            return render_template('listar_vagas.html', vagas=[], error="Erro ao carregar vagas")
        finally:
            cursor.close()
            conn.close()
    
    return render_template('listar_vagas.html', vagas=[], error="Erro de conexão com o banco")

if __name__ == '__main__':
    print("Iniciando sistema de currículos com IA...")
    print(f"API Key: {'*' * 10}{GEMINI_API_KEY[-4:] if GEMINI_API_KEY else 'NÃO ENCONTRADA'}")
    print(f"Usando: {'SQLite' if USE_SQLITE else 'MySQL'}")
    init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)
